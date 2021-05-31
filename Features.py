# import pip install docx2pdf
# import pip install python-docx
# import pip install xlsxwriter

import platform
import os


# Creats word file
def createWord(list, outputPath):
    import docx

    doc = docx.Document()
    doc.add_heading('            Auto Course Planner', 0)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    row = table.rows[0].cells
    row[0].text = 'Day'
    row[1].text = 'Hour'
    row[2].text = 'Class Type'
    row[3].text = 'Lesson Code'

    for i in list:
        row = table.add_row().cells
        row[0].text = i.day
        row[1].text = i.clock
        row[2].text = i.className
        row[3].text = i.code

    doc.save(outputPath + '/AutoCoursePlanner.docx')
    # os.system("AutoCoursePlanner.docx")   it opens the word file.

# createCSV
def createCSV(list,outputPath):
    import csv
    with open(outputPath+"/AutoCoursePlanner.csv", 'w', newline='') as file:
        writer = csv.writer(file, delimiter=';')
        writer.writerows([i.day,i.clock,i.className,i.code] for i in list)


# converst word file to pdf
def convertToPdf(source, destination):
    from docx2pdf import convert

    # source = ""
    # destination = ""
    convert(source, destination)


# creates excel file
def createExcell(list, outputPath):
    import xlsxwriter

    name = "/AutoCoursePlanner.xlsx"
    workbook = xlsxwriter.Workbook(outputPath + name)
    worksheet = workbook.add_worksheet()
    row = 1
    col = 0

    # className, day, clock, code
    worksheet.write(0, 0, "Day")
    worksheet.write(0, 1, "Clock")
    worksheet.write(0, 2, "Class Name")
    worksheet.write(0, 3, "Code")

    for item in (list):
        worksheet.write(row, col, item.day)
        worksheet.write(row, col + 1, item.clock)
        worksheet.write(row, col + 2, item.className)
        worksheet.write(row, col + 3, item.code)
        row += 1

    workbook.close()


# to install required packages
def packetInstaller():
    """to install required packages """
    packets = ["docx2pdf", "python-docx", "xlsxwriter", "tk", "csv"]

    if platform.system() == "Windows":
        os.system('cmd /c "pip install ' + " ".join(packets) + '"')
    else:
        os.system("pip install " + " ".join(packets))
